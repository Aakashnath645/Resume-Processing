import io
from fpdf import FPDF
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import logging

class EnhancedReport:
    """
    Generates enhanced reports in PDF and DOCX formats.
    """

    def __init__(self):
        """Initializes the EnhancedReport class."""
        pass

    def generate_detailed_report(self, analysis_results, report_type='pdf'):
        """
        Generates a detailed report in the specified format.

        Args:
            analysis_results (dict): A dictionary containing the analysis results.
            report_type (str, optional): The type of report to generate ('pdf' or 'docx'). Defaults to 'pdf'.

        Returns:
            bytes: The report in bytes format.
        """
        if report_type == 'pdf':
            return self._generate_pdf_report(analysis_results)
        elif report_type == 'docx':
            return self._generate_docx_report(analysis_results)
        else:
            raise ValueError("Invalid report type. Choose 'pdf' or 'docx'.")

    def _generate_pdf_report(self, analysis_results):
        """Generates a detailed PDF report with consistent formatting."""
        pdf = FPDF()
        pdf.add_page()
        
        # Add title with consistent formatting
        pdf.set_font("Arial", "B", size=16)
        pdf.cell(0, 10, txt="Resume Analysis Report", ln=1, align='C') # type: ignore
        pdf.ln(10)
        
        # Add candidate info with consistent sections
        pdf.set_font("Arial", "B", size=12)
        for field in ['Candidate Name', 'Department', 'Role']:
            if field in analysis_results:
                label = "Position: " if field == "Role" else f"{field}: "
                pdf.cell(0, 10, txt=f"{label}{analysis_results[field]}", ln=1) # type: ignore
        pdf.ln(5)

        # Match results section
        pdf.set_font("Arial", "B", size=12)
        pdf.cell(0, 10, txt="Analysis Results", ln=1) # type: ignore
        pdf.set_font("Arial", size=11)
        for field in ['Match %', 'Suitable']:
            if field in analysis_results:
                label = "Match Percentage: " if field == "Match %" else "Recommendation: "
                pdf.cell(0, 10, txt=f"{label}{analysis_results[field]}", ln=1) # type: ignore
        
        # Detailed scores section
        if "AI Scores" in analysis_results:
            pdf.ln(5)
            pdf.set_font("Arial", "B", size=12)
            pdf.cell(0, 10, txt="Detailed Scores", ln=1) # type: ignore
            pdf.set_font("Arial", size=11)
            for category, score in analysis_results["AI Scores"].items():
                pdf.cell(0, 10, txt=f"{category.title()}: {score}", ln=1) # type: ignore
        
        # Detailed analysis section
        if "Detailed Analysis" in analysis_results:
            pdf.ln(5)
            pdf.set_font("Arial", "B", size=12)
            pdf.cell(0, 10, txt="Detailed Analysis", ln=1) # type: ignore
            pdf.set_font("Arial", size=11)
            
            analysis_text = analysis_results["Detailed Analysis"]
            pdf.multi_cell(0, 10, txt=analysis_text) # type: ignore

        # Return bytes directly with fallback
        try:
            return bytes(pdf.output())
        except Exception:
            return pdf.output().encode('latin-1', errors='replace') # type: ignore

    def generate_batch_pdf_report(self, results):
        """Generate a combined PDF report for all results"""
        try:
            pdf = FPDF()
            pdf.add_page()
            
            # Add title
            pdf.set_font("Arial", "B", size=16)
            pdf.cell(0, 10, txt="Batch Resume Analysis Report", ln=1, align='C') # type: ignore
            pdf.ln(5)
            
            # Add summary
            pdf.set_font("Arial", size=12)
            pdf.cell(0, 10, txt=f"Total Candidates Analyzed: {len(results)}", ln=1) # type: ignore
            pdf.ln(10)
            
            # Process each candidate
            for i, result in enumerate(results, 1):
                pdf.set_font("Arial", "B", size=14)
                pdf.cell(0, 10, txt=f"Candidate Report {i}", ln=1) # type: ignore
                
                # Candidate details
                pdf.set_font("Arial", "B", size=12)
                for field in ['Candidate Name', 'Department', 'Role']:
                    if field in result:
                        label = "Position: " if field == "Role" else f"{field}: "
                        pdf.cell(0, 10, txt=f"{label}{result[field]}", ln=1) # type: ignore
                
                # Match results
                pdf.ln(5)
                for field in ['Match %', 'Suitable']:
                    if field in result:
                        label = "Match Percentage: " if field == "Match %" else "Recommendation: "
                        pdf.cell(0, 10, txt=f"{label}{result[field]}", ln=1) # type: ignore
                
                # Detailed scores
                if "AI Scores" in result:
                    pdf.ln(5)
                    pdf.set_font("Arial", "B", size=12)
                    pdf.cell(0, 10, txt="Detailed Scores", ln=1) # type: ignore
                    pdf.set_font("Arial", size=11)
                    for category, score in result["AI Scores"].items():
                        pdf.cell(0, 10, txt=f"{category.title()}: {score}", ln=1) # type: ignore
                
                # Detailed analysis
                if "Detailed Analysis" in result:
                    pdf.ln(5)
                    pdf.set_font("Arial", "B", size=12)
                    pdf.cell(0, 10, txt="Detailed Analysis", ln=1) # type: ignore
                    pdf.set_font("Arial", size=11)
                    pdf.multi_cell(0, 10, txt=result["Detailed Analysis"]) # type: ignore
                
                pdf.add_page()  # New page for next candidate
            
            # Return bytes directly with fallback
            try:
                return bytes(pdf.output())
            except Exception:
                return pdf.output().encode('latin-1', errors='replace') # type: ignore
                
        except Exception as e:
            logging.error(f"Error generating PDF report: {e}")
            return None

    def _generate_docx_report(self, analysis_results):
        """Generates a detailed DOCX report with consistent formatting."""
        document = Document()
        
        # Add title
        document.add_heading('Resume Analysis Report', level=1)
        
        # Add candidate info
        if "Candidate Name" in analysis_results:
            document.add_paragraph(f"Candidate: {analysis_results['Candidate Name']}")
        if "Department" in analysis_results:
            document.add_paragraph(f"Department: {analysis_results['Department']}")
        if "Role" in analysis_results:
            document.add_paragraph(f"Position: {analysis_results['Role']}")
        
        # Add scores section
        document.add_heading('Analysis Results', level=2)
        if "Match %" in analysis_results:
            document.add_paragraph(f"Match Percentage: {analysis_results['Match %']}")
        if "Suitable" in analysis_results:
            document.add_paragraph(f"Recommendation: {analysis_results['Suitable']}")
        
        # Add AI Scores if present
        if "AI Scores" in analysis_results:
            document.add_heading('Detailed Scores', level=2)
            for category, score in analysis_results["AI Scores"].items():
                document.add_paragraph(f"{category.title()}: {score}")
        
        # Add detailed analysis
        if "Detailed Analysis" in analysis_results:
            document.add_heading('Detailed Analysis', level=2)
            document.add_paragraph(analysis_results["Detailed Analysis"])

        docx_stream = io.BytesIO()
        document.save(docx_stream)
        docx_stream.seek(0)
        return docx_stream.read()

    def generate_batch_docx_report(self, batch_results):
        """Generates a comprehensive batch DOCX report."""
        document = Document()
        
        # Add title
        document.add_heading('Batch Resume Analysis Report', level=1)
        
        # Summary section
        document.add_paragraph(f"Total Candidates Analyzed: {len(batch_results)}")
        document.add_paragraph()
        
        # Individual reports
        for i, results in enumerate(batch_results, 1):
            document.add_heading(f'Candidate Report {i}', level=2)
            
            # Candidate details
            if "Candidate Name" in results:
                document.add_paragraph(f"Name: {results['Candidate Name']}")
            if "Department" in results:
                document.add_paragraph(f"Department: {results['Department']}")
            if "Role" in results:
                document.add_paragraph(f"Position: {results['Role']}")
            
            # Match results
            if "Match %" in results:
                document.add_paragraph(f"Match Percentage: {results['Match %']}")
            if "Suitable" in results:
                document.add_paragraph(f"Recommendation: {results['Suitable']}")
            
            # Detailed scores
            if "AI Scores" in results:
                document.add_heading('Detailed Scores', level=3)
                for category, score in results["AI Scores"].items():
                    document.add_paragraph(f"{category.title()}: {score}")
            
            # Detailed analysis
            if "Detailed Analysis" in results:
                document.add_heading('Detailed Analysis', level=3)
                document.add_paragraph(results["Detailed Analysis"])
            
            document.add_paragraph()  # Add space between reports
        
        docx_stream = io.BytesIO()
        document.save(docx_stream)
        docx_stream.seek(0)
        return docx_stream.read()
