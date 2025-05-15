import os
from dotenv import load_dotenv

# Wrap imports in try-except to handle missing packages
try:
    import streamlit as st
    import google.generativeai as genai
    import ollama
    import PyPDF2
    import docx
    from io import BytesIO
    from fpdf import FPDF
    from docx import Document
    import pandas as pd
    import json
    import sqlite3
    from datetime import datetime
    import logging
except ImportError as e:
    raise ImportError(f"Required package not found. Please run 'pip install streamlit python-dotenv google-generativeai ollama PyPDF2 python-docx fpdf2 pandas'. Error: {e}")

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)

# Validate and configure Gemini
if not os.getenv('GEMINI_API_KEY'):
    raise ValueError("GEMINI_API_KEY not found in environment variables. Please add it to .env file")

try:
    genai.configure(api_key=os.getenv('GEMINI_API_KEY')) # type: ignore
    model = genai.GenerativeModel('gemini-2.0-flash') # type: ignore
except Exception as e:
    raise RuntimeError(f"Failed to initialize Gemini model: {e}")

import streamlit as st
import ollama
import PyPDF2
import docx
from io import BytesIO
from fpdf import FPDF
from docx import Document
import time
import re
import streamlit.components.v1 as components
import sqlite3
from datetime import datetime
import pandas as pd
from collections import Counter
import json
from utils.ai_analysis import AIAnalyzer
from utils.report_generator import EnhancedReport
from typing import Dict, Any
import logging
import google.generativeai as genai
from utils.extract_utils import extract_candidate_name, validate_name

# Add these utility functions right after imports
def safe_score_calculation(score_value: float, default: float = 50.0) -> float:
    """
    Safely convert and validate a score value
    
    Args:
        score_value: The score to validate
        default: Default value if conversion fails
        
    Returns:
        float: Normalized score between 0 and 100
    """
    try:
        score = float(score_value)
        return min(100, max(0, score))
    except (TypeError, ValueError):
        return default

def validate_dict_structure(d, required_keys, default_values):
    """Ensure dictionary has required structure"""
    if not isinstance(d, dict):
        return default_values
    
    result = default_values.copy()
    for key in required_keys:
        if key in d and isinstance(d[key], type(default_values[key])):
            result[key] = d[key]
    return result

# Initialize components once at the top of the file
ai_analyzer = AIAnalyzer()
report_generator = EnhancedReport()

# Add after imports
def adapt_datetime(ts):
    return ts.isoformat()

def convert_datetime(val):
    try:
        return datetime.fromisoformat(val.decode())
    except AttributeError:
        return datetime.fromisoformat(val)

# Register adapters
sqlite3.register_adapter(datetime, adapt_datetime)
sqlite3.register_converter("timestamp", convert_datetime)

# Initialize database
def init_db():
    conn = sqlite3.connect('resume_analysis.db', detect_types=sqlite3.PARSE_DECLTYPES)
    c = conn.cursor()
    c.execute('''
        CREATE TABLE IF NOT EXISTS analyses (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            candidate_name TEXT,
            department TEXT,
            role TEXT,
            match_percentage REAL,
            suitable TEXT,
            detailed_analysis TEXT,
            timestamp timestamp,
            resume_file_name TEXT
        )
    ''')
    conn.commit()
    conn.close()

init_db()

# Add these database utility functions after imports
def get_db_connection():
    """Get database connection with proper timeout settings"""
    conn = sqlite3.connect('resume_analysis.db', 
                         detect_types=sqlite3.PARSE_DECLTYPES,
                         timeout=30,  # 30 second timeout
                         isolation_level=None)  # autocommit mode
    return conn

def safe_db_execute(query, params=None):
    """Execute database query with proper error handling"""
    max_retries = 3
    retry_delay = 1  # seconds
    
    for attempt in range(max_retries):
        try:
            conn = get_db_connection()
            try:
                with conn:  # auto-commits and handles rollback
                    cursor = conn.cursor()
                    if params:
                        cursor.execute(query, params)
                    else:
                        cursor.execute(query)
                    return True
            finally:
                conn.close()
        except sqlite3.OperationalError as e:
            if "database is locked" in str(e) and attempt < max_retries - 1:
                time.sleep(retry_delay)
                continue
            raise
        except Exception as e:
            st.error(f"Database error: {str(e)}")
            return False

def load_job_roles():
    with open("config/job_roles.json") as f:
        return json.load(f)

def process_resume( # type: ignore
    resume_text: str,
    job_description: str, 
    role_level: str,
    department: str = None # type: ignore
) -> Dict[str, Any]:
    """
    Process single resume with enhanced error handling
    
    Args:
        resume_text: The text content of the resume
        job_description: The job description text
        role_level: The role level (e.g. junior, senior)
        department: Optional - department name (reserved for future use)
        
    Returns:
        Dict containing score, AI analysis, ATS score and detailed analysis
    """
    default_structure = {
        'score': 50,
        'ai_analysis': {
            'scores': {
                'overall': 50,
                'detailed': {
                    'technical': 50,
                    'experience': 50,
                    'leadership': 50,
                    'cultural': 50
                }
            }
        },
        'ats_score': 50,
        'detailed_analysis': 'Analysis not available'
    }

    try:
        # Get keyword score as fallback
        keyword_score = compute_basic_keyword_score(resume_text, job_description)
        
        # Get AI analysis with validation
        ai_results = ai_analyzer.parallel_analysis(resume_text, job_description, role_level) # type: ignore
        ai_results = validate_dict_structure(
            ai_results,
            ['scores', 'explanation'],
            {'scores': {'overall': 50, 'detailed': {}}, 'explanation': 'No analysis available'}
        )
        
        # Get detailed analysis
        detailed_analysis = analyze_resume_with_jd(resume_text, job_description, job_description)
        if not detailed_analysis or len(detailed_analysis) < 10:
            detailed_analysis = "Could not generate detailed analysis"

        # Calculate final score
        ai_score = safe_score_calculation(ai_results['scores'].get('overall', 50))
        final_score = (ai_score * 0.6) + (keyword_score * 0.4)
        
        return {
            'score': safe_score_calculation(final_score),
            'ai_analysis': ai_results,
            'ats_score': keyword_score,
            'detailed_analysis': detailed_analysis
        }
    except Exception as e:
        logging.error(f"Error in resume processing: {str(e)}")
        # Use keyword score for basic functionality
        basic_score = compute_basic_keyword_score(resume_text, job_description)
        return {
            'score': safe_score_calculation(basic_score),
            'ai_analysis': default_structure['ai_analysis'],
            'ats_score': basic_score,
            'detailed_analysis': f"Error in analysis: {str(e)}"
        }

# ----------------------------
# Helper Functions
# ----------------------------

def sanitize_filename(name):
    sanitized = re.sub(r'[^\w\-_ ]', '', name)
    return sanitized.replace(" ", "_")

def extract_text_from_resume(uploaded_file):
    text = ""
    if uploaded_file:
        ext = uploaded_file.name.split(".")[-1].lower()
        try:
            if ext == "pdf":
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                for page in pdf_reader.pages:
                    t = page.extract_text()
                    if t:
                        text += t + "\n"
            elif ext == "docx":
                doc = Document(BytesIO(uploaded_file.read()))
                for para in doc.paragraphs:
                    text += para.text + "\n"
            elif ext == "doc":
                st.warning(f"Note: .doc file detected. Convert to .docx for best results.")
                text = uploaded_file.read().decode('utf-8', errors='ignore')
            else:
                text = uploaded_file.read().decode("utf-8")
            return text.strip()
        except Exception as e:
            st.error(f"Error processing file {uploaded_file.name}: {str(e)}")
            return f"Error processing file: {str(e)}"
    return text.strip()

def extract_text_from_job_description_file(uploaded_file):
    text = ""
    if uploaded_file:
        ext = uploaded_file.name.split(".")[-1].lower()
        if ext == "pdf":
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            for page in pdf_reader.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
        elif ext in ["doc", "docx"]:
            doc_obj = docx.Document(uploaded_file)
            for para in doc_obj.paragraphs:
                text += para.text + "\n"
        else:
            text = uploaded_file.read().decode("utf-8")
    return text.strip()

# --- ATS Score Functions ---
def compute_ATS_score_keyword(resume_text, job_description):
    """
    Compute ATS score based on frequency of job description keywords in the resume.
    Returns a score between 0 and 100.
    """
    job_words = re.findall(r'\w+', job_description.lower())
    resume_words = re.findall(r'\w+', resume_text.lower())
    job_counter = Counter(job_words)
    resume_counter = Counter(resume_words)
    common_score = 0
    for word, freq in job_counter.items():
        if word in resume_counter:
            common_score += min(resume_counter[word], freq)
    max_score = sum(job_counter.values())
    score_percentage = (common_score / max_score) * 100 if max_score > 0 else 0
    return min(score_percentage, 100)

def compute_ATS_score_combined(resume_text, job_description, role_level):
    """Enhanced scoring with AI and basic keyword matching"""
    try:
        # Get AI analysis
        ai_results = ai_analyzer.parallel_analysis(resume_text, job_description, role_level) # type: ignore
        
        # Basic keyword matching
        keyword_score = compute_basic_keyword_score(resume_text, job_description)
        
        # Calculate final score
        final_score = (
            ai_results['scores']['overall'] * 0.6 +  # AI analysis
            keyword_score * 0.4                      # Basic keyword matching
        )
        
        return {
            'score': min(final_score, 100),
            'ai_analysis': ai_results,
            'keyword_score': keyword_score
        }
    except Exception as e:
        print(f"Error in scoring: {e}")
        # Fallback to basic scoring
        return {
            'score': compute_basic_keyword_score(resume_text, job_description),
            'ai_analysis': {'scores': {'overall': 50}, 'explanation': 'AI analysis failed'},
            'keyword_score': 50
        }

def compute_basic_keyword_score(resume_text, job_description):
    """Modified keyword matching with more lenient scoring"""
    job_words = set(job_description.lower().split())
    resume_words = set(resume_text.lower().split())
    
    # Calculate match percentage with bonus for partial matches
    common_words = job_words.intersection(resume_words)
    partial_matches = sum(1 for job_word in job_words 
                         if any(job_word in resume_word for resume_word in resume_words))
    
    if not job_words:
        return 0
    
    exact_match_score = (len(common_words) / len(job_words)) * 100
    partial_match_bonus = (partial_matches / len(job_words)) * 20  # Bonus for partial matches
    
    return min(100, exact_match_score + partial_match_bonus)

def get_ats_keyword_analysis(job_description, resume_text):
    """Enhanced ATS keyword analysis with context understanding"""
    # Clean and prepare texts
    job_text = preprocess_text(job_description)
    resume_text = preprocess_text(resume_text)
    
    # Extract key phrases and skills
    required_skills = extract_key_phrases(job_text)
    resume_skills = extract_key_phrases(resume_text)
    
    # Calculate different types of matches
    exact_matches = calculate_exact_matches(required_skills, resume_skills)
    semantic_matches = calculate_semantic_matches(required_skills, resume_skills)
    context_matches = analyze_skill_context(required_skills, resume_text)
    
    return {
        'keyword_match': (exact_matches * 0.6 + semantic_matches * 0.25 + context_matches * 0.15),
        'missing_critical': list(identify_missing_critical_skills(required_skills, resume_skills)),
        'skill_relevance': calculate_skill_relevance_scores(required_skills, resume_skills)
    }

def get_comprehensive_ai_analysis(resume_text, job_description, role_level):
    """Comprehensive AI analysis using multiple prompts for accuracy"""
    try:
        # Technical Skills Analysis
        tech_prompt = create_technical_analysis_prompt(resume_text, job_description, role_level)
        tech_score = get_ai_response_with_validation(tech_prompt)
        
        # Leadership and Management Analysis
        leadership_prompt = create_leadership_analysis_prompt(resume_text, role_level)
        leadership_score = get_ai_response_with_validation(leadership_prompt)
        
        # Project and Impact Analysis
        project_prompt = create_project_analysis_prompt(resume_text, role_level)
        project_score = get_ai_response_with_validation(project_prompt)
        
        # Soft Skills and Cultural Fit Analysis
        cultural_prompt = create_cultural_fit_prompt(resume_text, job_description)
        cultural_score = get_ai_response_with_validation(cultural_prompt)
        
        return {
            'technical_skills': normalize_score(tech_score.get('overall', 0)),
            'role_specific_skills': normalize_score(tech_score.get('role_specific', 0)),
            'leadership': normalize_score(leadership_score.get('leadership', 0)),
            'management_skills': normalize_score(leadership_score.get('management', 0)),
            'strategic_thinking': normalize_score(leadership_score.get('strategic', 0)),
            'project_scale': normalize_score(project_score.get('scale', 0)),
            'impact_score': normalize_score(project_score.get('impact', 0)),
            'complexity_handling': normalize_score(project_score.get('complexity', 0)),
            'communication': normalize_score(cultural_score.get('communication', 0)),
            'teamwork': normalize_score(cultural_score.get('teamwork', 0)),
            'adaptability': normalize_score(cultural_score.get('adaptability', 0))
        }
    except Exception as e:
        print(f"Error in comprehensive analysis: {e}")
        return {
            'technical_skills': 0, 'role_specific_skills': 0,
            'leadership': 0, 'management_skills': 0,
            'strategic_thinking': 0, 'project_scale': 0,
            'impact_score': 0, 'complexity_handling': 0,
            'communication': 0, 'teamwork': 0, 'adaptability': 0
        }

def get_role_specific_criteria(role_level):
    """Return role-specific evaluation criteria based on role level"""
    role_lower = role_level.lower()
    
    if any(term in role_lower for term in ['senior', 'lead', 'manager', 'director', 'chief']):
        return """
        - Strategic thinking and vision
        - Leadership and mentoring experience
        - Architecture and system design expertise
        - Track record of successful project delivery
        - Team management and growth
        """
    elif any(term in role_lower for term in ['junior', 'intern', 'trainee', 'associate']):
        return """
        - Foundational technical knowledge
        - Learning ability and enthusiasm
        - Basic coding and problem-solving skills
        - Team collaboration potential
        - Growth mindset
        """
    else:  # Mid-level
        return """
        - Strong technical implementation skills
        - Project ownership and delivery
        - Team collaboration and mentoring
        - Technical decision-making ability
        - Balance of hands-on and leadership skills
        """

def create_technical_analysis_prompt(resume_text, job_description, role_level):
    return f"""Carefully evaluate the technical capabilities for this {role_level} position.

Consider:
1. Technical Skill Match
- Core required skills: {extract_technical_skills(job_description)}
- Skill proficiency level
- Hands-on vs. oversight experience

2. Project Complexity
- Scale and scope of projects
- Technical challenges handled
- Architecture and design decisions

3. Level-Appropriate Analysis:
{get_role_specific_criteria(role_level)}

Resume Text: {resume_text[:1000]}
Job Description: {job_description[:1000]}

Provide balanced scoring with detailed justification."""

def create_leadership_analysis_prompt(resume_text, role_level):
    return f"Analyze the leadership and management skills in the resume for the role of {role_level}. Resume: {resume_text}"

def create_project_analysis_prompt(resume_text, role_level):
    return f"Analyze the project and impact details in the resume for the role of {role_level}. Resume: {resume_text}"

def create_cultural_fit_prompt(resume_text, job_description):
    return f"Analyze the soft skills and cultural fit in the resume for the job described as: {job_description}. Resume: {resume_text}"

def normalize_score(score):
    try:
        return min(100, max(0, float(score)))
    except (ValueError, TypeError):
        return 0

def calculate_experience_score(years: float, role_level: str) -> float:
    """
    Calculate experience score based on years and role level
    
    Args:
        years: Number of years of experience
        role_level: Level of the role (e.g. junior, senior)
        
    Returns:
        float: Score between 0 and 100
    """
    role_lower = role_level.lower()
    
    # Senior/Leadership roles
    if any(term in role_lower for term in ['senior', 'lead', 'manager', 'director', 'chief']):
        if years >= 8: return 100
        if years >= 5: return 80
        if years >= 3: return 60
        return max(30, years * 10)
        
    # Junior roles
    elif any(term in role_lower for term in ['junior', 'intern', 'trainee', 'associate']):
        if years >= 3: return 100
        if years >= 1: return 80
        return max(40, years * 40)
        
    # Mid-level roles
    else:
        if years >= 5: return 100
        if years >= 3: return 80
        if years >= 1: return 60
        return max(30, years * 20)

def analyze_experience_depth(resume_text, role_level):
    """Detailed analysis of experience quality and relevance"""
    exp_prompt = f"""Analyze the professional experience in detail:
1. Total years of relevant experience
2. Career progression and growth
3. Role-specific experience relevance
4. Project complexity progression
5. Leadership responsibility growth

Resume: {resume_text[:2000]}"""

    try:
        response = get_ai_response_with_validation(exp_prompt)
        years = extract_years_from_response(response['content'])
        
        return {
            'years': years,
            'years_score': calculate_experience_score(years, role_level),
            'relevance_score': response.get('relevance', 70),
            'progression_score': response.get('progression', 70)
        }
    except Exception as e:
        print(f"Experience analysis error: {e}")
        return {'years': 0, 'years_score': 50, 'relevance_score': 50, 'progression_score': 50}

def calculate_role_weights(role_level):
    """Calculate importance weights based on role level"""
    weights = {
        'Chief': {
            'skills': 0.15,
            'experience': 0.25,
            'leadership': 0.30,
            'project': 0.15,
            'cultural': 0.15
        },
        'Director': {
            'skills': 0.20,
            'experience': 0.25,
            'leadership': 0.25,
            'project': 0.15,
            'cultural': 0.15
        },
        'Senior': {
            'skills': 0.25,
            'experience': 0.20,
            'leadership': 0.20,
            'project': 0.20,
            'cultural': 0.15
        },
        'Manager': {
            'skills': 0.25,
            'experience': 0.20,
            'leadership': 0.20,
            'project': 0.20,
            'cultural': 0.15
        },
        'Associate': {
            'skills': 0.30,
            'experience': 0.15,
            'leadership': 0.15,
            'project': 0.20,
            'cultural': 0.20
        }
    }
    
    for level, weight in weights.items():
        if level in role_level:
            return weight
    return weights['Associate']

def get_ai_response_with_validation(prompt: str, retries: int = 3) -> Dict[str, Any]:
    """Get AI response with validation and error handling"""
    for attempt in range(retries):
        try:
            response = ollama.chat(model="gemma:2b", messages=[{
                "role": "user",
                "content": prompt
            }])
            
            result = parse_and_validate_ai_response(response)
            if result:
                return result
        except Exception as e:
            logging.error(f"AI analysis attempt {attempt + 1} failed: {e}")
            if attempt == retries - 1:
                return {'content': '', 'scores': {}}
        time.sleep(1)
    return {'content': '', 'scores': {}}

def extract_scores_from_content(content):
    """Extract numerical scores from AI response content"""
    scores = {}
    try:
        # Look for patterns like "Score: 85" or "Technical Score: 90"
        score_patterns = [
            r'(\w+)\s*score:\s*(\d+)',
            r'score\s*[:-]\s*(\d+)',
            r'rating\s*[:-]\s*(\d+)',
            r'(\d+)\s*\/\s*100'
        ]
        
        for pattern in score_patterns:
            matches = re.findall(pattern, content.lower())
            for match in matches:
                if isinstance(match, tuple):
                    category, score = match
                    scores[category] = float(score)
                else:
                    scores['overall'] = float(match)
        
        # Ensure there's at least an overall score
        if not scores:
            scores['overall'] = 50  # Default score
            
    except Exception as e:
        print(f"Error extracting scores: {e}")
        scores['overall'] = 50
        
    return scores

def parse_and_validate_ai_response(response):
    """Parse and validate AI response with score normalization"""
    try:
        content = response.get("message", {}).get("content", "").strip()
        scores = extract_scores_from_content(content)
        return {
            'content': content,
            'scores': normalize_scores(scores)
        }
    except:
        return None

def is_suitable_numeric(score, role_level, threshold=45):  # Lowered base threshold
    """Determine suitability with role-based adjustments"""
    role_lower = role_level.lower()
    
    # Experienced roles
    if any(term in role_lower for term in ['senior', 'lead', 'manager', 'director', 'chief']):
        return "Yes" if score >= 40 else "No"  # More lenient for senior roles
    
    # Junior/Entry roles
    elif any(term in role_lower for term in ['junior', 'intern', 'trainee', 'associate']):
        return "Yes" if score >= 35 else "No"  # Even more lenient for junior roles
    
    # Mid-level roles
    return "Yes" if score >= threshold else "No"

def analyze_job_description(job_description):
    return job_description

def analyze_resume_with_jd(resume_text: str, job_description: str, role_level: str) -> str: # type: ignore
    """Generate objective analysis using Gemini 2.0 Flash"""
    try:
        prompt = f"""Provide an objective evaluation of this candidate's qualifications:

        Role Context:
        Position: {role_level}
        Requirements: {job_description[:2000]}

        Candidate Profile:
        {resume_text[:2000]}

        Evaluation Guidelines:
        1. Assess qualifications objectively:
           - Core competencies match
           - Relevant accomplishments
           - Applicable experience
           
        2. Consider diverse experience types:
           - Direct industry experience
           - Transferable skills
           - Cross-domain expertise
           - Project achievements
           
        3. Focus on role-specific capabilities:
           - Technical/Professional skills
           - Domain knowledge
           - Demonstrated outcomes
           
        4. Provide balanced assessment:
           - Key strengths
           - Areas for growth
           - Overall role alignment
           
        Important: Base evaluation solely on role-relevant qualifications. 
        Exclude any bias related to demographics, background, or non-essential factors.
        
        Format: Clear, structured analysis with specific supporting examples."""

        if ai_analyzer is None:
            return "AI analysis unavailable. Using basic matching only."

        response = ai_analyzer.gemini.generate_analysis(prompt)
        return response['text'] if response['success'] else "Analysis could not be generated."
    except Exception as e:
        logging.error(f"Gemini API error: {str(e)}")
        return f"Could not generate analysis. Please check your API key configuration."

def format_match_percentage(match_percentage):
    try:
        return f"{float(match_percentage):.0f}%"
    except:
        return "0%"

def load_custom_ui():
    with open("style.css") as f:
        st.markdown(f"<style>{f.read()}</style>", unsafe_allow_html=True)
    with open("animations.html") as f:
        components.html(f.read(), height=0)

# Add delay for AI response
def get_ai_response(prompt, max_retries=3, delay=1):
    """Helper function to handle AI response with retries"""
    for attempt in range(max_retries):
        try:
            response = genai.generate_content(prompt) # type: ignore
            return response.text.strip()
        except Exception as e:
            if attempt < max_retries - 1:
                time.sleep(delay)
                continue
            print(f"AI Error after {max_retries} attempts: {e}")
            return ""
    return ""

st.set_page_config(
    page_title="Resume Screening Assistant",
    page_icon="🔍",
    layout="wide"
)

load_custom_ui()

st.markdown('<div class="main-container">', unsafe_allow_html=True)

# Header
st.markdown('''
    <div class="header">
        <h1>🔍 Smart Resume Screening Assistant</h1>
        <p class="subtitle">Upload resumes and match them with job descriptions</p>
    </div>
''', unsafe_allow_html=True)

# Department Section
st.markdown("**Department**")
col1, col2 = st.columns([3, 1])

job_description = st.session_state.get("job_description_text", "")

# Add this before the UI code
def reset_all():
    """Reset all session state and generate new keys"""
    if st.button("Clear Upload(s)", key="clear_uploads_btn"):
        clear_uploads()
        st.rerun()  # Only rerun after explicit clear

def clear_uploads():
    """Clear only upload-related session state"""
    keys_to_clear = [
        'job_description_text',
        'uploaded_files',
        'reset_counter',
        'results'  # Add results to clear
    ]
    for key in keys_to_clear:
        if key in st.session_state:
            del st.session_state[key]
    st.success("Uploads and results cleared successfully!")

# Modify the file upload sections with dynamic keys
st.markdown("### Job Description")
jd_tab1, jd_tab2 = st.tabs(["📝 Text Input", "📎 Upload File"])

# Get job description first
with jd_tab1:
    job_description = st.text_area(
        "Enter Job Description",
        value="",
        height=150,
        key=f"jd_text_input_{st.session_state.get('reset_counter', 0)}"
    )

with jd_tab2:
    jd_file = st.file_uploader(
        "Upload Job Description",
        type=['pdf','docx','txt'],
        key=f"jd_file_upload_{st.session_state.get('reset_counter', 0)}",
        accept_multiple_files=False
    )
    if jd_file is not None:
        job_description = extract_text_from_job_description_file(jd_file)
        st.success("Job description uploaded successfully")

# Store job description in session state
if job_description:
    st.session_state.job_description_text = job_description

# Department and Role Selection
# Update the department selection section in main UI
st.markdown("### Department and Role Selection")
col1, col2 = st.columns(2)
with col1:
    selected_department = st.selectbox(
        "Select Department",
        options=list(load_job_roles().keys()),
        key="department_select"
    )

with col2:
    selected_role = st.selectbox(
        "Select Role",
        options=load_job_roles()[selected_department]["roles"],
        key="role_select"
    )

st.session_state.current_department = selected_department
st.session_state.current_role = selected_role

# Resume Upload Section with dynamic key
st.markdown("### Resume Upload")
uploaded_files = st.file_uploader(
    label="Upload Resume(s)",
    type=['pdf','docx','doc'],
    accept_multiple_files=True,
    key=f"resume_files_{st.session_state.get('reset_counter', 0)}",
    help="Upload one or more resumes"
)

# Modify the buttons section to remove duplicate "Start New Analysis" buttons
# Modified Clear and Action Buttons (combine into single row)
col1, col2 = st.columns([1,1])
with col1:
    process_button = st.button("Process Resume(s)", type="primary", key="process_resumes_btn")
with col2:
    reset_all()  # Call reset function instead of direct button

# Dashboard placeholders
if 'dashboard_containers' not in st.session_state:
    st.session_state.dashboard_containers = {
        'progress': st.empty(),
        'status': st.empty(),
        'results': st.empty()
    }

# Add these loading message lists
loading_messages = [
    "🧠 AI is analyzing the resume...",
    "📊 Calculating match scores...",
    "🔍 Evaluating technical skills...",
    "📈 Assessing experience level...",
    "🎯 Determining role fit...",
    "🤝 Analyzing cultural match...",
    "📑 Generating detailed report...",
    "✨ Finalizing analysis..."
]

fun_facts = [
    "Did you know? The first resume was written by Leonardo da Vinci in 1482!",
    "On average, recruiters spend only 7.4 seconds reviewing a resume initially.",
    "75% of large companies use ATS to screen resumes.",
    "The most overused word in resumes is 'experienced'.",
    "Including a photo on your resume can lead to 88% rejection rate.",
    "40% of hiring managers spend less than a minute reviewing a resume."
]

# Initialize session state for results tracking
if 'current_batch_results' not in st.session_state:
    st.session_state.current_batch_results = []

# Process Resumes
if process_button:
    if uploaded_files and job_description:
        progress_container = st.session_state.dashboard_containers['progress']
        status_container = st.session_state.dashboard_containers['status']
        results_container = st.session_state.dashboard_containers['results']
        
        # Clear previous results
        st.session_state.current_batch_results = []
        
        # Process job description
        status_container.markdown("""
            <div class="analysis-status">
                <h3>🎯 Starting Resume Analysis</h3>
                <p>Preparing to analyze resumes against job requirements...</p>
            </div>
        """, unsafe_allow_html=True)
        
        parsed_jd = analyze_job_description(job_description)
        results = []
        total_files = len(uploaded_files)
        
        # Process each resume with enhanced progress display
        for idx, file in enumerate(uploaded_files, 1):
            # Update progress with random messages and fun facts
            progress = idx / total_files
            current_message = loading_messages[idx % len(loading_messages)]
            current_fact = fun_facts[idx % len(fun_facts)]
            
            progress_container.markdown(f"""
                <div class="progress-container">
                    <div class="progress-bar" style="width: {progress*100}%"></div>
                    <div class="progress-status">
                        <h4>{current_message}</h4>
                        <p>Processing: {file.name}</p>
                        <div class="fun-fact">💡 {current_fact}</div>
                    </div>
                </div>
            """, unsafe_allow_html=True)
            
            # Process resume
            text = extract_text_from_resume(file)
            if not text:
                st.error(f"Could not extract text from {file.name}")
                continue
                
            # Use new name extraction with validation
            name = extract_candidate_name(text)
            if name == "Unknown_Candidate":
                # Use filename as fallback, cleaned up
                base_name = file.name.rsplit('.', 1)[0]
                name = f"Candidate_{base_name}"
            
            # Use the manually selected department and role
            department_to_use = st.session_state.current_department
            role_to_use = st.session_state.current_role
            
            # Process resume
            analysis = process_resume(
                text, 
                job_description, 
                st.session_state.current_role,
                st.session_state.current_department
            )
            
            # Ensure we have valid scores
            match_percentage = analysis.get('score', 0)
            suitable = is_suitable_numeric(match_percentage, st.session_state.current_role)
            
            # Store in database with safe execution
            query = '''
                INSERT INTO analyses 
                (candidate_name, department, role, match_percentage, suitable, 
                 detailed_analysis, ai_scores, ats_score, resume_file_name)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            '''
            params = (
                name, 
                st.session_state.current_department,
                st.session_state.current_role,
                match_percentage,
                suitable,
                analysis.get('detailed_analysis', ''),
                json.dumps(analysis.get('ai_analysis', {}).get('scores', {})),
                analysis.get('ats_score', 0),
                file.name
            )
            
            if not safe_db_execute(query, params):
                st.error(f"Failed to save analysis for {name}")
                continue
            
            # Append to current batch instead of session state
            st.session_state.current_batch_results.append({
                "Candidate Name": name,
                "Department": st.session_state.current_department,
                "Role": st.session_state.current_role,
                "Match %": f"{analysis['score']:.0f}%",
                "Suitable": is_suitable_numeric(analysis['score'], st.session_state.current_role),
                "Detailed Analysis": analysis['detailed_analysis'],
                "AI Scores": analysis['ai_analysis']['scores']['detailed'],
                "ATS Score": analysis['ats_score']
            })
            
            time.sleep(0.5) 
        progress_container.empty()
        status_container.empty()
            

if st.session_state.current_batch_results:
    st.success(f"✨ Analysis completed! Processed {len(st.session_state.current_batch_results)} resumes.")

    # Generate batch reports with error handling
    try:
        pdf_batch_data = report_generator.generate_batch_pdf_report(st.session_state.current_batch_results)
        docx_batch_data = report_generator.generate_batch_docx_report(st.session_state.current_batch_results)
    except Exception as e:
        st.error(f"Error generating reports: {str(e)}")
        pdf_batch_data = None
        docx_batch_data = None

    # Single instance of download buttons
    col1, col2 = st.columns(2)
    current_time = int(time.time())

    with col1:
        if pdf_batch_data is not None:
            st.download_button(
                "📄 Download All PDF Reports",
                data=pdf_batch_data,
                file_name=f"batch_analysis_reports_{current_time}.pdf",
                mime="application/pdf",
                key=f"pdf_download_{current_time}",
                use_container_width=True
            )
        else:
            st.error("PDF report generation failed")

    with col2:
        if docx_batch_data is not None:
            st.download_button(
                "📑 Download All DOCX Reports",
                data=docx_batch_data,
                file_name=f"batch_analysis_reports_{current_time}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key=f"docx_download_{current_time}",
                use_container_width=True
            )
        else:
            st.error("DOCX report generation failed")

    # Updated results card display with consistent naming
    st.markdown("### Analysis Results")
    for i, res in enumerate(st.session_state.current_batch_results):
        # Clean and format candidate name
        candidate_name = res.get('Candidate Name', 'Unknown Candidate')
        if candidate_name.startswith('Candidate_'):
            # Remove 'Candidate_' prefix and replace underscores with spaces
            candidate_name = ' '.join(candidate_name.split('_')[1:]).title()
        
        with st.container():
            st.markdown(f"""
                <div class="result-card">
                    <div class="result-header">
                        <h3>{candidate_name}</h3>
                        <div class="match-badge">{res['Match %']}</div>
                    </div>
                    <div class="result-content">
                        <p><strong>Department:</strong> {res['Department']}</p>
                        <p><strong>Role:</strong> {res['Role']}</p>
                        <p><strong>Recommendation:</strong> 
                            <span class="{'success' if res['Suitable'] == 'Yes' else 'warning'}">
                                {res['Suitable']}
                            </span>
                        </p>
                    </div>
                </div>
            """, unsafe_allow_html=True)
            
            with st.expander(f"View Detailed Analysis for {candidate_name}"):
                st.write(res['Detailed Analysis'])

st.markdown('</div>', unsafe_allow_html=True)

# Initialize AI analyzer
ai_analyzer = AIAnalyzer()
report_generator = EnhancedReport()

# Add these functions to replace removed text_processing dependencies
def preprocess_text(text):
    """Basic text preprocessing without spacy"""
    text = text.lower()
    text = re.sub(r'[^\w\s]', ' ', text)
    return ' ' .join(text.split())

def extract_key_phrases(text):
    """Simple key phrase extraction without spacy"""
    words = text.lower().split()
    # Basic skill patterns
    skill_patterns = {
        'programming': r'\b(python|java|javascript|c\+\+|ruby|php)\b',
        'tools': r'\b(git|docker|kubernetes|jenkins|aws|azure)\b',
        'soft_skills': r'\b(leadership|communication|teamwork|management)\b'
    }
    
    phrases = set()
    for category, pattern in skill_patterns.items():
        matches = re.findall(pattern, text.lower())
        phrases.update(matches)
    
    return phrases

def normalize_scores(scores):
    """Normalize scores to 0-100 range"""
    if isinstance(scores, (int, float)):
        return min(100, max(0, float(scores)))
    return 50

def calculate_exact_matches(required, provided):
    """Calculate exact keyword matches"""
    if not required:
        return 0
    matches = len(required.intersection(provided))
    return (matches / len(required)) * 100

def calculate_semantic_matches(required, provided):
    """Calculate semantic matches between required and provided skills"""
    score = 0
    if not required:
        return 0
        
    for req_skill in required:
        # Check for partial matches
        for prov_skill in provided:
            if req_skill in prov_skill or prov_skill in req_skill:
                score += 1
                break
                
    return (score / len(required)) * 100

def analyze_skill_context(required_skills, resume_text):
    """Simple skill context analysis"""
    context_score = 0
    sentences = resume_text.lower().split('.')
    
    for skill in required_skills:
        for sentence in sentences:
            if skill in sentence:
                if any(exp in sentence for exp in ['years', 'experience', 'expert']):
                    context_score += 1
    
    return min(100, context_score * 10)

def calculate_skill_relevance_scores(required, provided):
    """Calculate relevance scores for skills"""
    return {skill: 100 if skill in provided else 0 for skill in required}

def identify_missing_critical_skills(required, provided):
    """Identify missing skills"""
    return required - provided

# Initialize components at the top
ai_analyzer = AIAnalyzer()
report_generator = EnhancedReport()

def process_resume(resume_text, job_description, role_level, department=None): # type: ignore
    """Process single resume with enhanced error handling"""
    default_structure = {
        'score': 50,
        'ai_analysis': {
            'scores': {
                'overall': 50,
                'detailed': {
                    'technical': 50,
                    'experience': 50,
                    'leadership': 50,
                    'cultural': 50
                }
            }
        },
        'ats_score': 50,
        'detailed_analysis': 'Analysis not available'
    }

    try:
        # Role-specific scoring adjustments
        is_senior = any(term in role_level.lower() for term in ['senior', 'lead', 'manager', 'director', 'chief'])
        is_junior = any(term in role_level.lower() for term in ['junior', 'intern', 'trainee', 'associate'])
        
        # Get base scores
        keyword_score = compute_basic_keyword_score(resume_text, job_description)
        ai_results = ai_analyzer.parallel_analysis(resume_text, job_description, role_level) # type: ignore
        
        # Apply role-based weights
        if is_senior:
            weights = {
                'ai_score': 0.4,
                'keyword_score': 0.2,
                'experience_bonus': 0.4
            }
            exp_years = extract_years_from_response(ai_results.get('content', ''))
            exp_bonus = min(100, exp_years * 10)  # 10 points per year, max 100
        elif is_junior:
            weights = {
                'ai_score': 0.3,
                'keyword_score': 0.4,
                'potential_bonus': 0.3
            }
            # Consider education and trainings more for junior roles
            potential_score = analyze_potential(resume_text, job_description)
        else:
            weights = {
                'ai_score': 0.4,
                'keyword_score': 0.3,
                'balanced_bonus': 0.3
            }
            balanced_score = analyze_balanced_skills(resume_text, job_description)
        
        # Calculate final score with role-specific adjustments
        ai_score = safe_score_calculation(ai_results['scores'].get('overall', 50))
        if is_senior:
            final_score = (
                ai_score * weights['ai_score'] +
                keyword_score * weights['keyword_score'] +
                exp_bonus * weights['experience_bonus']
            )
        elif is_junior:
            final_score = (
                ai_score * weights['ai_score'] +
                keyword_score * weights['keyword_score'] +
                potential_score * weights['potential_bonus']
            )
        else:
            final_score = (
                ai_score * weights['ai_score'] +
                keyword_score * weights['keyword_score'] +
                balanced_score * weights['balanced_bonus']
            )
        
        # Generate detailed analysis
        detailed_analysis = analyze_resume_with_jd(resume_text, job_description, role_level)
        if not detailed_analysis or len(detailed_analysis) < 10:
            detailed_analysis = "Could not generate detailed analysis"
        
        return {
            'score': safe_score_calculation(final_score),
            'ai_analysis': ai_results,
            'ats_score': keyword_score,
            'detailed_analysis': detailed_analysis
        }
        
    except Exception as e:
        logging.error(f"Error in resume processing: {str(e)}")
        return default_structure

def analyze_potential(resume_text, job_description):
    """Analyze potential for junior roles"""
    score = 50  # Base score
    
    # Check for relevant education
    education_keywords = ['degree', 'bachelor', 'master', 'certification', 'training']
    if any(keyword in resume_text.lower() for keyword in education_keywords):
        score += 15
    
    # Check for projects/internships
    project_keywords = ['project', 'internship', 'trainee', 'volunteer']
    if any(keyword in resume_text.lower() for keyword in project_keywords):
        score += 15
    
    # Check for technical skills match
    tech_skills = extract_technical_skills(job_description)
    if any(skill in resume_text.lower() for skill in tech_skills):
        score += 20
    
    return min(100, score)

def extract_years_from_response(text):
    """Extract years of experience from text"""
    # Look for common experience patterns
    patterns = [
        r'(\d+)\+?\s*years?\s*(of)?\s*experience',
        r'(\d+)\+?\s*years?\s*(in|at)',
        r'worked\s*(\d+)\+?\s*years?',
        r'experience\s*(\d+)\+?\s*years?'
    ]
    
    total_years = 0
    for pattern in patterns:
        matches = re.findall(pattern, text.lower())
        if matches:
            # Take the largest year number found
            years = max([int(match[0]) for match in matches])
            total_years = max(total_years, years)
    
    return total_years

def analyze_balanced_skills(resume_text, job_description):
    """Analyze balanced skill set for mid-level roles"""
    score = 50  # Base score
    
    # Check experience (2-5 years typical for mid-level)
    years = extract_years_from_response(resume_text)
    if 2 <= years <= 5:
        score += 20
    elif years > 5:
        score += 15  # Slightly lower for overqualified
    
    # Check for project leadership
    if any(term in resume_text.lower() for term in ['lead', 'managed', 'coordinated']):
        score += 15
    
    # Check for relevant skills
    tech_skills = extract_technical_skills(job_description)
    skill_match = len([skill for skill in tech_skills if skill in resume_text.lower()])
    score += min(15, skill_match * 3)  # Up to 15 points for skills
    
    return min(100, score)

def extract_technical_skills(text):
    """Extract technical skills from text"""
    common_tech_skills = [
        'python', 'java', 'javascript', 'sql', 'aws', 'azure',
        'docker', 'kubernetes', 'react', 'angular', 'node',
        'machine learning', 'ai', 'data science', 'devops'
    ]
    return [skill for skill in common_tech_skills if skill in text.lower()]

def process_resume(resume_text: str, job_description: str, role_level: str, department: str = None) -> Dict[str, Any]: # type: ignore
    """Enhanced resume processing with better AI-score integration"""
    try:
        # Get AI analysis first
        ai_results = ai_analyzer.parallel_analysis(resume_text, job_description, role_level) # type: ignore
        
        # Extract detailed scores
        technical_score = ai_results.get('scores', {}).get('detailed', {}).get('technical', 50)
        experience_score = ai_results.get('scores', {}).get('detailed', {}).get('experience', 50)
        keyword_score = compute_basic_keyword_score(resume_text, job_description)
        
        # Role-specific scoring adjustments
        is_senior = any(term in role_level.lower() for term in ['senior', 'lead', 'manager', 'director'])
        is_junior = any(term in role_level.lower() for term in ['junior', 'intern', 'trainee'])
        
        # Calculate weighted scores based on role
        if is_senior:
            weights = {
                'technical': 0.35,
                'experience': 0.40,
                'keyword': 0.25
            }
            # Apply minimum thresholds for senior roles
            if experience_score < 60:
                experience_score *= 0.8
            if technical_score < 65:
                technical_score *= 0.8
        elif is_junior:
            weights = {
                'technical': 0.45,
                'experience': 0.15,
                'keyword': 0.40
            }
            # More emphasis on potential for junior roles
            if technical_score > 70:
                technical_score *= 1.2
        else:  # Mid-level
            weights = {
                'technical': 0.40,
                'experience': 0.35,
                'keyword': 0.25
            }
        
        # Calculate final score
        final_score = (
            technical_score * weights['technical'] +
            experience_score * weights['experience'] +
            keyword_score * weights['keyword']
        )
        
        # Generate detailed analysis using AI
        analysis_text = analyze_resume_with_jd(resume_text, job_description, role_level)
        
        return {
            'score': min(100, max(0, final_score)),
            'ai_analysis': {
                'scores': {
                    'overall': final_score,
                    'detailed': {
                        'technical': technical_score,
                        'experience': experience_score,
                        'keyword': keyword_score
                    }
                }
            },
            'ats_score': keyword_score,
            'detailed_analysis': analysis_text
        }
        
    except Exception as e:
        logging.error(f"Resume processing error: {str(e)}")
        return {
            'score': 50,
            'ai_analysis': {'scores': {'overall': 50, 'detailed': {}}},
            'ats_score': 50,
            'detailed_analysis': f'Error in analysis: {str(e)}'
        }

def analyze_resume_with_jd(resume_text: str, job_description: str, role_level: str) -> str: # type: ignore
    """Generate AI analysis with improved scoring justification"""
    try:
        prompt = f"""Analyze this resume for a {role_level} position with detailed scoring:

        Job Requirements: {job_description[:1000]}
        Resume: {resume_text[:1000]}

        Please provide:
        1. Technical Skills Score (0-100):
           - Required skills match
           - Skill proficiency level
           - Relevant technical experience
           
        2. Experience Score (0-100):
           - Years of relevant experience
           - Project complexity
           - Role-specific achievements
           
        3. Overall Match Analysis:
           - Key strengths
           - Areas for improvement
           - Specific examples from resume
           
        4. Final Recommendation:
           - Scoring justification
           - Role suitability
           - Key differentiators

        Format: Clear sections with numeric scores and detailed explanations."""

        response = model.generate_content(prompt)
        return response.text if response else "Analysis could not be generated."
    except Exception as e:
        logging.error(f"AI analysis error: {str(e)}")
        return "Could not generate detailed analysis. Using basic matching only."

# Process uploaded files
if process_button and uploaded_files and job_description:
    for idx, file in enumerate(uploaded_files, 1):
        text = extract_text_from_resume(file)
        if not text:
            st.error(f"Could not extract text from {file.name}")
            continue

        # Use new name extraction with validation
        name = extract_candidate_name(text)
        if name == "Unknown_Candidate":
            # Use filename as fallback, cleaned up
            base_name = file.name.rsplit('.', 1)[0]
            name = f"Candidate_{base_name}"
        
        # Get analysis with proper role and department
        analysis = process_resume(
            text,
            job_description,
            st.session_state.current_role,
            st.session_state.current_department
        )
        
        # Ensure we have valid scores
        match_percentage = analysis.get('score', 0)
        suitable = is_suitable_numeric(match_percentage, st.session_state.current_role)
        
        # Store in database
        query = '''
            INSERT INTO analyses 
            (candidate_name, department, role, match_percentage, suitable, 
             detailed_analysis, ai_scores, ats_score, resume_file_name)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        '''
        params = (
            name,
            st.session_state.current_department,
            st.session_state.current_role,
            match_percentage,
            suitable,
            analysis.get('detailed_analysis', ''),
            json.dumps(analysis.get('ai_analysis', {}).get('scores', {})),
            analysis.get('ats_score', 0),
            file.name
        )
        
        if not safe_db_execute(query, params):
            st.error(f"Failed to save analysis for {name}")
            continue

# Wrap AI analyzer initialization in try-except
try:
    ai_analyzer = AIAnalyzer()
except ValueError as e:
    st.error(f"Error initializing AI: {str(e)}")
    st.warning("System will use basic keyword matching for analysis")
    ai_analyzer = None

def analyze_resume_with_jd(resume_text: str, job_description: str, role_level: str) -> str:
    """Generate analysis with fallback to basic analysis"""
    if ai_analyzer is None:
        # Fallback to basic keyword analysis
        keyword_score = compute_basic_keyword_score(resume_text, job_description)
        return f"""Basic Analysis (AI Unavailable):
        - Keyword Match Score: {keyword_score:.0f}%
        - Recommendation based on keyword matching only."""
        
    try:
        prompt = f"""Analyze resume briefly for {role_level} position.
        Requirements: {job_description[:1000]}
        Resume: {resume_text[:1000]}"""

        response = model.generate_content(prompt)
        return response.text if response else "Analysis could not be generated."
    except Exception as e:
        logging.error(f"Analysis error: {str(e)}")
        return "Could not generate detailed analysis. Using basic matching only."